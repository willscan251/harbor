"""
Harbor - File Watcher
Monitors an inbox folder and automatically sorts files using AI analysis.

Supports routing to BOTH client folders and company-level folders:
  - Clients/[Client Name]/[Category]
  - Company/[Finance|Legal|HR|Admin]
  - Marketing/[Brand Assets|Website|Social Media|Proposals]
  - Resources/[Templates|Training|Reference]

SETUP:
1. Create a folder on your desktop called "Harbor Inbox"
2. Set HARBOR_INBOX=/path/to/folder in .env
3. For SharePoint upload: python integrations/sharepoint.py auth
4. Run: python integrations/file_watcher.py start

Files dropped in the inbox will be:
1. Text extracted from the file (PDF, DOCX, XLSX, TXT, images)
2. Analyzed by AI to determine destination folder
3. Uploaded to SharePoint (if configured) OR moved to local folder
4. Moved to local _Processed folder (or _unsorted if truly unclassifiable)
"""

import os
import sys
import time
import shutil
import logging
from datetime import datetime
from pathlib import Path
from threading import Thread

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

import config
import database as db

# Try importing watchdog
try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    WATCHDOG_AVAILABLE = True
except ImportError:
    WATCHDOG_AVAILABLE = False
    print("⚠ watchdog not installed. Run: pip install watchdog")

# Try importing AI processor
try:
    import ai_processor
    AI_AVAILABLE = config.ANTHROPIC_API_KEY is not None
except ImportError:
    AI_AVAILABLE = False
    print("⚠ ai_processor not available")

# Try importing SharePoint client
try:
    from integrations.sharepoint import SharePointClient
    SHAREPOINT_AVAILABLE = True
except ImportError:
    try:
        from sharepoint import SharePointClient
        SHAREPOINT_AVAILABLE = True
    except ImportError:
        SHAREPOINT_AVAILABLE = False

# Document text extraction libraries
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠ PyMuPDF not installed. Run: pip install pymupdf")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠ python-docx not installed. Run: pip install python-docx")

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("⚠ openpyxl not installed. Run: pip install openpyxl")

logger = logging.getLogger('integration.file_watcher')

# File types we'll process
ALLOWED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
    '.txt', '.csv', '.png', '.jpg', '.jpeg', '.gif', '.zip',
    '.vtt', '.srt', '.md', '.rtf'
}

# Minimum file age before processing (seconds) - prevents processing partial uploads
MIN_FILE_AGE = 2


# ============================================
# Text Extraction
# ============================================

def extract_text(file_path: str, max_chars: int = 3000) -> str:
    """
    Extract text content from a file for AI analysis.
    Supports: PDF, DOCX, XLSX, TXT, CSV, MD, VTT, SRT
    
    Returns extracted text or None if extraction fails.
    """
    ext = Path(file_path).suffix.lower()
    
    try:
        # Plain text files
        if ext in ['.txt', '.md', '.csv', '.vtt', '.srt', '.rtf']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read(max_chars)
        
        # PDF files
        elif ext == '.pdf' and PDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                text = ""
                # Read up to first 5 pages for better context
                for page_num in range(min(5, len(doc))):
                    page = doc[page_num]
                    text += page.get_text()
                    if len(text) >= max_chars:
                        break
                doc.close()
                return text[:max_chars] if text.strip() else None
            except Exception as e:
                logger.warning(f"PDF extraction failed for {file_path}: {e}")
                return None
        
        # Word documents
        elif ext == '.docx' and DOCX_AVAILABLE:
            try:
                doc = DocxDocument(file_path)
                text = ""
                # Get paragraph text
                for para in doc.paragraphs:
                    text += para.text + "\n"
                    if len(text) >= max_chars:
                        break
                # Also check tables for content
                if len(text) < max_chars:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                            if len(text) >= max_chars:
                                break
                        if len(text) >= max_chars:
                            break
                return text[:max_chars] if text.strip() else None
            except Exception as e:
                logger.warning(f"DOCX extraction failed for {file_path}: {e}")
                return None
        
        # Excel files
        elif ext in ['.xlsx', '.xls'] and XLSX_AVAILABLE:
            try:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                text = ""
                # Read first 3 sheets, 50 rows each
                for sheet_name in wb.sheetnames[:3]:
                    ws = wb[sheet_name]
                    text += f"Sheet: {sheet_name}\n"
                    row_count = 0
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c) if c is not None else "" for c in row]
                        text += " | ".join(cells) + "\n"
                        row_count += 1
                        if row_count >= 50 or len(text) >= max_chars:
                            break
                    if len(text) >= max_chars:
                        break
                wb.close()
                return text[:max_chars] if text.strip() else None
            except Exception as e:
                logger.warning(f"XLSX extraction failed for {file_path}: {e}")
                return None
        
        # Images - flag for vision-based analysis
        elif ext in ['.png', '.jpg', '.jpeg', '.gif']:
            return "__IMAGE_FILE__"
            
    except Exception as e:
        logger.warning(f"Text extraction failed for {file_path}: {e}")
    
    return None


# ============================================
# File Watcher
# ============================================

class FileWatcherIntegration:
    """File system watcher that auto-sorts files to client and company folders"""
    
    def __init__(self):
        self.inbox = config.FILE_INBOX
        self.client_files_root = config.CLIENT_FILES_ROOT
        self.observer = None
        self.is_running = False
        self.sharepoint = None
        
        # Try to initialize SharePoint client
        if SHAREPOINT_AVAILABLE:
            try:
                self.sharepoint = SharePointClient()
                if self.sharepoint.load_token():
                    logger.info("SharePoint client loaded successfully")
                else:
                    logger.info("SharePoint not authenticated - using local storage only")
                    self.sharepoint = None
            except Exception as e:
                logger.warning(f"SharePoint init failed: {e}")
                self.sharepoint = None
    
    def is_configured(self) -> bool:
        """Check if file watcher is configured"""
        return bool(self.inbox)
    
    def test_connection(self) -> tuple:
        """Test file watcher setup. Returns (success, message)"""
        if not self.inbox:
            return False, "HARBOR_INBOX not configured in .env"
        
        if not os.path.exists(self.inbox):
            try:
                os.makedirs(self.inbox)
                return True, f"Created inbox folder: {self.inbox}"
            except Exception as e:
                return False, f"Cannot create inbox folder: {e}"
        
        file_count = len([f for f in os.listdir(self.inbox) 
                          if not f.startswith('.') and not f.startswith('_')])
        
        sp_status = "Connected" if self.sharepoint else "Not configured (local only)"
        ai_status = "Available" if AI_AVAILABLE else "Not available"
        
        return True, (
            f"Inbox: {self.inbox} ({file_count} files)\n"
            f"  SharePoint: {sp_status}\n"
            f"  AI categorization: {ai_status}\n"
            f"  PDF extraction: {'Yes' if PDF_AVAILABLE else 'No - pip install pymupdf'}\n"
            f"  DOCX extraction: {'Yes' if DOCX_AVAILABLE else 'No - pip install python-docx'}\n"
            f"  XLSX extraction: {'Yes' if XLSX_AVAILABLE else 'No - pip install openpyxl'}\n"
            f"  Image vision: {'Yes' if AI_AVAILABLE else 'No - needs ANTHROPIC_API_KEY'}"
        )
    
    def process_file(self, file_path: str) -> dict:
        """
        Process a single file from the inbox.
        
        Returns: {success, destination, client, category, confidence, reason}
        """
        filename = os.path.basename(file_path)
        ext = Path(file_path).suffix.lower()
        
        # Skip unsupported file types
        if ext not in ALLOWED_EXTENSIONS:
            return {'success': False, 'reason': f'Unsupported file type: {ext}'}
        
        # Skip if file is too new (might still be uploading)
        try:
            file_age = time.time() - os.path.getmtime(file_path)
            if file_age < MIN_FILE_AGE:
                return {'success': False, 'reason': 'File still uploading'}
        except:
            pass
        
        logger.info(f"Processing: {filename}")
        print(f"\n📄 Processing: {filename}")
        
        # Step 1: Extract text from file
        content_preview = None
        is_image = False
        if AI_AVAILABLE:
            try:
                content_preview = extract_text(file_path)
                if content_preview == "__IMAGE_FILE__":
                    is_image = True
                    content_preview = None
                    logger.info(f"  Image file detected - will use vision analysis")
                    print(f"  🖼️  Image file - using AI vision to read content")
                elif content_preview:
                    logger.info(f"  Extracted {len(content_preview)} chars of text")
                    print(f"  📖 Extracted {len(content_preview)} chars of text")
                else:
                    logger.info(f"  No text could be extracted from this file")
                    print(f"  ⚠ Cannot extract text from this file type - sending to _unsorted")
            except Exception as e:
                logger.warning(f"  Text extraction failed: {e}")
        
        # Step 2: AI categorization
        # RULE: Never sort based on filename alone. Must have content or vision.
        analysis = None
        if AI_AVAILABLE:
            if is_image:
                # Use Claude's vision API to read the image
                try:
                    analysis = ai_processor.categorize_image_file(filename, file_path)
                    logger.info(f"  Vision AI result: {analysis}")
                    print(f"  🤖 AI Vision: {analysis.get('destination', '?')} (confidence: {analysis.get('confidence', '?')})")
                    print(f"     Reason: {analysis.get('reason', 'N/A')}")
                except Exception as e:
                    logger.warning(f"  AI vision analysis failed: {e}")
                    print(f"  ❌ AI vision analysis failed: {e}")
            elif content_preview:
                # Use text-based categorization
                try:
                    analysis = ai_processor.categorize_file(filename, content_preview)
                    logger.info(f"  AI result: {analysis}")
                    print(f"  🤖 AI: {analysis.get('destination', '?')} (confidence: {analysis.get('confidence', '?')})")
                    print(f"     Reason: {analysis.get('reason', 'N/A')}")
                except Exception as e:
                    logger.warning(f"  AI analysis failed: {e}")
                    print(f"  ❌ AI analysis failed: {e}")
            else:
                # No content extracted and not an image — cannot make an informed decision
                print(f"  🚫 No content available - will not sort based on filename alone")
                return self._move_to_unsorted(file_path, filename, {
                    'confidence': 'none',
                    'reason': f'Cannot extract content from {ext} file - needs manual review'
                })
        
        # Step 3: Route the file
        if analysis and analysis.get('confidence') in ['high', 'medium']:
            destination = analysis.get('destination', '_unsorted')
            
            # Validate destination path makes sense
            if destination and destination != '_unsorted':
                return self._route_file(file_path, filename, ext, analysis)
        
        # Step 4: If low confidence or no analysis, use _unsorted
        return self._move_to_unsorted(file_path, filename, analysis)
    
    def _route_file(self, file_path: str, filename: str, ext: str, analysis: dict) -> dict:
        """Route a file to its determined destination"""
        destination = analysis.get('destination', '_unsorted')
        client_id = analysis.get('client_id')
        
        # Upload to SharePoint if available
        if self.sharepoint:
            try:
                sp_path = f"{destination}/{filename}"
                result = self.sharepoint.upload_file(file_path, sp_path)
                if result:
                    logger.info(f"  Uploaded to SharePoint: {sp_path}")
                    print(f"  ☁️  Uploaded to SharePoint: {sp_path}")
                    
                    # Move local file to _Processed
                    processed_folder = os.path.join(self.inbox, '_Processed')
                    os.makedirs(processed_folder, exist_ok=True)
                    dest_local = os.path.join(processed_folder, filename)
                    if os.path.exists(dest_local):
                        base, extension = os.path.splitext(filename)
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        dest_local = os.path.join(processed_folder, f"{base}_{timestamp}{extension}")
                    shutil.move(file_path, dest_local)
                    
                    # Log to database
                    self._log_sort(filename, ext, client_id, destination, analysis)
                    
                    return {
                        'success': True,
                        'destination': f"SharePoint: {sp_path}",
                        'client': analysis.get('client_code'),
                        'category': analysis.get('category'),
                        'confidence': analysis.get('confidence'),
                        'reason': analysis.get('reason', '')
                    }
            except Exception as e:
                logger.warning(f"  SharePoint upload failed: {e}, falling back to local")
                print(f"  ⚠ SharePoint upload failed: {e}")
        
        # Fall back to local folder storage
        local_dest_folder = os.path.join(self.client_files_root, destination)
        os.makedirs(local_dest_folder, exist_ok=True)
        
        dest_path = os.path.join(local_dest_folder, filename)
        if os.path.exists(dest_path):
            base, extension = os.path.splitext(filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            dest_path = os.path.join(local_dest_folder, f"{base}_{timestamp}{extension}")
        
        shutil.move(file_path, dest_path)
        
        # Log to database
        self._log_sort(filename, ext, client_id, destination, analysis)
        
        logger.info(f"  Sorted to: {destination}/")
        print(f"  ✅ Sorted to: {destination}/")
        
        return {
            'success': True,
            'destination': dest_path,
            'client': analysis.get('client_code'),
            'category': analysis.get('category'),
            'confidence': analysis.get('confidence'),
            'reason': analysis.get('reason', '')
        }
    
    def _move_to_unsorted(self, file_path: str, filename: str, analysis: dict = None) -> dict:
        """Move file to _unsorted as a last resort"""
        unsorted_folder = os.path.join(self.client_files_root, '_unsorted')
        os.makedirs(unsorted_folder, exist_ok=True)
        
        dest_path = os.path.join(unsorted_folder, filename)
        if os.path.exists(dest_path):
            base, ext = os.path.splitext(filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            dest_path = os.path.join(unsorted_folder, f"{base}_{timestamp}{ext}")
        
        shutil.move(file_path, dest_path)
        
        reason = analysis.get('reason', 'Could not determine destination') if analysis else 'AI not available'
        logger.info(f"  Moved to _unsorted/: {reason}")
        print(f"  📂 Moved to _unsorted/: {reason}")
        
        return {
            'success': True,
            'destination': dest_path,
            'client': None,
            'category': '_unsorted',
            'confidence': analysis.get('confidence') if analysis else 'none',
            'reason': reason
        }
    
    def _log_sort(self, filename: str, ext: str, client_id: int, destination: str, analysis: dict):
        """Log the file sort to the database"""
        try:
            # Add to documents table if it's a client file
            if client_id:
                db.execute_db('''
                    INSERT INTO documents (client_id, filename, display_name, file_path, file_type, category, client_visible, uploaded_by)
                    VALUES (?, ?, ?, ?, ?, ?, 0, 'file_watcher')
                ''', [
                    client_id,
                    filename,
                    filename,
                    destination,
                    ext[1:] if ext else None,
                    analysis.get('category', 'general')
                ])
            
            # Log activity
            db.log_activity(
                action='document_sorted',
                description=f"Auto-sorted '{filename}' → {destination} ({analysis.get('confidence', '?')} confidence: {analysis.get('reason', '')})",
                client_id=client_id,
                entity_type='document',
                performed_by='file_watcher'
            )
        except Exception as e:
            logger.warning(f"Failed to log sort: {e}")
    
    def process_existing_files(self) -> int:
        """Process any files already in the inbox. Returns count processed."""
        if not os.path.exists(self.inbox):
            return 0
        
        files = [f for f in os.listdir(self.inbox)
                 if os.path.isfile(os.path.join(self.inbox, f))
                 and not f.startswith('.')
                 and not f.startswith('~')]
        
        if not files:
            return 0
        
        logger.info(f"Found {len(files)} existing file(s) in inbox")
        print(f"\n📥 Found {len(files)} file(s) in inbox to process")
        processed = 0
        
        for filename in files:
            file_path = os.path.join(self.inbox, filename)
            try:
                result = self.process_file(file_path)
                if result.get('success'):
                    processed += 1
            except Exception as e:
                logger.error(f"Error processing {filename}: {e}")
                print(f"  ❌ Error processing {filename}: {e}")
        
        print(f"\n✅ Processed {processed}/{len(files)} files")
        return processed
    
    def start(self):
        """Start watching the inbox folder"""
        if not WATCHDOG_AVAILABLE:
            raise RuntimeError("watchdog not installed. Run: pip install watchdog")
        
        # Create inbox if needed
        os.makedirs(self.inbox, exist_ok=True)
        os.makedirs(self.client_files_root, exist_ok=True)
        
        # Process existing files
        self.process_existing_files()
        
        # Set up event handler
        handler = InboxHandler(self)
        
        self.observer = Observer()
        self.observer.schedule(handler, self.inbox, recursive=False)
        self.observer.start()
        self.is_running = True
        
        logger.info(f"File watcher started. Watching: {self.inbox}")
    
    def stop(self):
        """Stop watching"""
        if self.observer:
            self.observer.stop()
            self.observer.join()
            self.is_running = False
            logger.info("File watcher stopped")
    
    def run_forever(self):
        """Start watching and block until interrupted"""
        self.start()
        
        print("\n" + "=" * 60)
        print("  Harbor - File Watcher")
        print("=" * 60)
        print(f"  Watching:    {self.inbox}")
        print(f"  Local files: {self.client_files_root}")
        print(f"  SharePoint:  {'Connected' if self.sharepoint else 'Not configured'}")
        print(f"  AI:          {'Enabled' if AI_AVAILABLE else 'Disabled'}")
        print(f"  PDF reader:  {'Yes' if PDF_AVAILABLE else 'No'}")
        print(f"  DOCX reader: {'Yes' if DOCX_AVAILABLE else 'No'}")
        print(f"  XLSX reader: {'Yes' if XLSX_AVAILABLE else 'No'}")
        print()
        print("  Drop files into the inbox folder to auto-sort them!")
        print("  Press Ctrl+C to stop")
        print("=" * 60)
        
        try:
            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            print("\n\nStopping...")
        finally:
            self.stop()


class InboxHandler(FileSystemEventHandler):
    """Handle file system events in the inbox folder"""
    
    def __init__(self, watcher: FileWatcherIntegration):
        self.watcher = watcher
        self.processing = set()
    
    def on_created(self, event):
        if event.is_directory:
            return
        
        file_path = event.src_path
        filename = os.path.basename(file_path)
        
        # Skip hidden, temp files, and _Processed folder
        if filename.startswith('.') or filename.startswith('~'):
            return
        if '_Processed' in file_path:
            return
        
        # Avoid double-processing
        if file_path in self.processing:
            return
        
        self.processing.add(file_path)
        
        # Wait for file to finish writing
        time.sleep(MIN_FILE_AGE + 0.5)
        
        try:
            if os.path.exists(file_path):
                self.watcher.process_file(file_path)
        finally:
            self.processing.discard(file_path)
    
    def on_modified(self, event):
        # Also handle modifications (some apps create then modify)
        self.on_created(event)


# Global instance
file_watcher = FileWatcherIntegration()


# ============================================
# CLI
# ============================================

if __name__ == '__main__':
    import sys
    
    def print_usage():
        print("""
Harbor - File Watcher CLI

Commands:
    start               Start watching inbox folder (runs forever)
    test                Test configuration and show capabilities
    process <path>      Process a single file or folder
    reprocess           Re-process files in _unsorted
    status              Show current status
    
Examples:
    python integrations/file_watcher.py start
    python integrations/file_watcher.py test
    python integrations/file_watcher.py process ~/Downloads/report.pdf
    python integrations/file_watcher.py reprocess
""")
    
    if len(sys.argv) < 2:
        # Default: start watching
        if not WATCHDOG_AVAILABLE:
            print("Error: watchdog not installed")
            print("Run: pip install watchdog")
            sys.exit(1)
        
        file_watcher.run_forever()
        sys.exit(0)
    
    cmd = sys.argv[1].lower()
    
    if cmd == 'start':
        if not WATCHDOG_AVAILABLE:
            print("Error: watchdog not installed")
            print("Run: pip install watchdog")
            sys.exit(1)
        file_watcher.run_forever()
    
    elif cmd == 'test':
        success, msg = file_watcher.test_connection()
        print(f"{'✓' if success else '✗'} File Watcher Status:")
        print(f"  {msg}")
    
    elif cmd == 'process' and len(sys.argv) > 2:
        path = sys.argv[2]
        if os.path.isfile(path):
            result = file_watcher.process_file(path)
            print(f"\nResult: {result}")
        elif os.path.isdir(path):
            count = 0
            for filename in sorted(os.listdir(path)):
                file_path = os.path.join(path, filename)
                if os.path.isfile(file_path) and not filename.startswith('.'):
                    result = file_watcher.process_file(file_path)
                    count += 1
            print(f"\nProcessed {count} files")
        else:
            print(f"Path not found: {path}")
    
    elif cmd == 'reprocess':
        unsorted_path = os.path.join(config.CLIENT_FILES_ROOT, '_unsorted')
        if os.path.exists(unsorted_path):
            files = [f for f in os.listdir(unsorted_path) 
                     if os.path.isfile(os.path.join(unsorted_path, f)) 
                     and not f.startswith('.')]
            if files:
                print(f"\n📂 Re-processing {len(files)} files from _unsorted/")
                for filename in sorted(files):
                    file_path = os.path.join(unsorted_path, filename)
                    result = file_watcher.process_file(file_path)
            else:
                print("No files in _unsorted/")
        else:
            print("_unsorted folder doesn't exist")
    
    elif cmd == 'status':
        print(f"\nFile Watcher Status")
        print(f"  Inbox: {config.FILE_INBOX}")
        print(f"  Client files: {config.CLIENT_FILES_ROOT}")
        print(f"  Watchdog: {'Available' if WATCHDOG_AVAILABLE else 'Not installed'}")
        print(f"  AI: {'Available' if AI_AVAILABLE else 'Not available'}")
        print(f"  SharePoint: {'Available' if SHAREPOINT_AVAILABLE else 'Not available'}")
        print(f"  PDF extraction: {'Yes' if PDF_AVAILABLE else 'No'}")
        print(f"  DOCX extraction: {'Yes' if DOCX_AVAILABLE else 'No'}")
        print(f"  XLSX extraction: {'Yes' if XLSX_AVAILABLE else 'No'}")
        
        if os.path.exists(config.FILE_INBOX):
            files = [f for f in os.listdir(config.FILE_INBOX) if not f.startswith('.') and not f.startswith('_')]
            print(f"  Files in inbox: {len(files)}")
        
        unsorted = os.path.join(config.CLIENT_FILES_ROOT, '_unsorted')
        if os.path.exists(unsorted):
            files = [f for f in os.listdir(unsorted) if not f.startswith('.')]
            print(f"  Files in _unsorted: {len(files)}")
    
    else:
        print(f"Unknown command: {cmd}")
        print_usage()
